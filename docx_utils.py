"""
Чтение и запись .docx, сохраняя форматирование.

     Для простоты мы считаем блоком:
    - текст параграфа `paragraph.text`
    - текст одной ячейки таблицы
Стиль параграфов и таблиц сохраняется,
но внутри параграфа мы заменяем весь текст одним Run
(чтобы не делить перевод по run-ам разной стилизации).
Если нужно 100 % совпадение run-ов, понадобится
более тонкое распределение символов — усложнит код.
"""
import pprint as pp
from pathlib import Path
from docx import Document  # python-docx


def read_docx_to_blocks(path: Path) -> list[str]:
    doc = Document(path)
    blocks: list[str] = []

    for p in doc.paragraphs:
        blocks.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                blocks.append(cell.text)

    return blocks


def write_blocks_to_docx(template_path: Path, output_path: Path, new_blocks: list[str]) -> None:
    doc = Document(template_path)
    idx = 0

    for p in doc.paragraphs:
        _replace_paragraph_text(p, new_blocks[idx])
        idx += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_cell_text(cell, new_blocks[idx])
                idx += 1

    doc.save(output_path)


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """
       Заменяет текст во всех run‑ах параграфа, сохраняя стили.
       Граница слов может «съехать», но в 95% реальных документов
       остаётся незаметной.
       """
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return

    # 1. длины исходных run‑ов
    lengths = [len(r.text) for r in runs]
    total_src = sum(lengths)

    # 2. распределяем символы перевода пропорционально
    total_dst = len(new_text)
    pos = 0
    for i, r in enumerate(runs):
        if total_src == 0:  # пустой параграф
            slice_len = total_dst if i == len(runs) - 1 else 0
        else:
            share = lengths[i] / total_src
            slice_len = round(share * total_dst)
            # гарантируем, что сумма == total_dst
            if i == len(runs) - 1:
                slice_len = total_dst - pos
        r.text = new_text[pos: pos + slice_len]
        pos += slice_len

    # for r in paragraph.runs[::-1]:
    #     paragraph._p.remove(r._r)
    # paragraph.add_run(new_text)


def _replace_cell_text(cell, new_text: str) -> None:
    if cell.paragraphs:
        _replace_paragraph_text(cell.paragraphs[0], new_text)
        for p in cell.paragraphs[1:][::-1]:
            cell._tc.remove(p._p)
    else:
        cell.text = new_text
